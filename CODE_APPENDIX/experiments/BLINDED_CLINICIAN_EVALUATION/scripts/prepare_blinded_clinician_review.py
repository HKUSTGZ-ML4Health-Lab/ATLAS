#!/usr/bin/env python3
# AAAI-27 paper reference: Blinded Clinician Evaluation and Table 2.
"""Prepare a blinded clinician-review packet for ATLAS vs Gemini.

The script samples an equal number of cases from the Western benchmark and
GeriMedBench. It never reads gold labels. It randomizes case order and A/B
assignment independently for each reviewer, removes model identifiers, and
exports reviewer packets, rating templates, a sample manifest, and a private
randomization key.
"""

from __future__ import annotations

import argparse
import csv
import json
import random
import re
import sys
from copy import deepcopy
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

MODEL_ATLAS = "ATLAS"
MODEL_GEMINI = "Gemini 3.1 Pro Preview"

BLOCKED_KEY_PATTERN = re.compile(
    r"(^|_)(gold|reference|label|labels|answer[_-]?key|target|expected|evaluator|score|metric)(_|$)",
    re.IGNORECASE,
)

CRITERIA = [
    ("Clinical correctness", "The medication decisions match the full patient state and the relevant guideline evidence."),
    ("Medication safety", "The output avoids harmful options and states the required cautions or monitoring."),
    ("Decision completeness", "The output covers recommendation, avoidance, caution or monitoring, and safer alternatives when relevant."),
    ("Actionability", "The output gives clear steps that a clinician or pharmacist can use in medication review."),
    ("Evidence consistency", "The medication claims agree with the evidence path and the patient facts shown in the packet."),
]


def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(description=__doc__)
    p.add_argument("--western-cases", required=True)
    p.add_argument("--western-atlas", required=True)
    p.add_argument("--western-gemini", required=True)
    p.add_argument("--geri-cases", required=True)
    p.add_argument("--geri-atlas", required=True)
    p.add_argument("--geri-gemini", required=True)
    p.add_argument("--out-dir", required=True)
    p.add_argument("--reviewers", type=int, default=3)
    p.add_argument("--n-per-benchmark", type=int, default=20)
    p.add_argument("--seed", type=int, default=20260729)
    p.add_argument(
        "--case-id-key",
        default="case_id",
        help="Primary case identifier key. The script also checks id and uid.",
    )
    return p.parse_args()


def load_records(path: str | Path) -> list[dict[str, Any]]:
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(p)
    text = p.read_text(encoding="utf-8-sig").strip()
    if not text:
        return []
    if p.suffix.lower() == ".jsonl":
        return [json.loads(line) for line in text.splitlines() if line.strip()]
    obj = json.loads(text)
    if isinstance(obj, list):
        return obj
    if isinstance(obj, dict):
        for key in ("cases", "predictions", "data", "items", "records"):
            if isinstance(obj.get(key), list):
                return obj[key]
    raise ValueError(f"Unsupported JSON structure in {p}")


def get_case_id(obj: dict[str, Any], preferred: str = "case_id") -> str:
    for key in (preferred, "case_id", "id", "uid"):
        value = obj.get(key)
        if value is not None and str(value).strip():
            return str(value)
    raise KeyError(f"No case identifier found in keys: {list(obj)[:20]}")


def index_records(records: Iterable[dict[str, Any]], preferred: str) -> dict[str, dict[str, Any]]:
    out: dict[str, dict[str, Any]] = {}
    for item in records:
        cid = get_case_id(item, preferred)
        if cid in out:
            raise ValueError(f"Duplicate case_id: {cid}")
        out[cid] = item
    return out


def sanitize_case(value: Any, parent_key: str = "") -> Any:
    """Remove reference labels and evaluator-only fields from patient data."""
    if isinstance(value, dict):
        result: dict[str, Any] = {}
        for key, val in value.items():
            key_s = str(key)
            if BLOCKED_KEY_PATTERN.search(key_s):
                continue
            if key_s.lower() in {
                "final_decision",
                "prediction",
                "model_output",
                "trace_verification",
                "strict_success",
                "osrs",
            }:
                continue
            result[key_s] = sanitize_case(val, key_s)
        return result
    if isinstance(value, list):
        return [sanitize_case(v, parent_key) for v in value]
    return value


def find_final_decision(pred: dict[str, Any]) -> dict[str, Any]:
    for key in ("final_decision", "decision", "output", "prediction"):
        val = pred.get(key)
        if isinstance(val, dict):
            if key in {"output", "prediction"} and isinstance(val.get("final_decision"), dict):
                return val["final_decision"]
            return val
    return pred


def as_text(value: Any) -> str:
    if value is None:
        return "Not stated"
    if isinstance(value, bool):
        return "Yes" if value else "No"
    if isinstance(value, str):
        return value.strip() or "Not stated"
    if isinstance(value, (int, float)):
        return str(value)
    if isinstance(value, list):
        if not value:
            return "None"
        return "\n".join(f"• {as_text(v)}" for v in value)
    if isinstance(value, dict):
        if not value:
            return "None"
        return "\n".join(f"{k}: {as_text(v)}" for k, v in value.items())
    return str(value)


def first_present(d: dict[str, Any], keys: Iterable[str]) -> Any:
    for key in keys:
        if key in d:
            return d[key]
    return None


def normalized_output(pred: dict[str, Any]) -> list[tuple[str, str]]:
    """Render both methods through the same field order without rewriting content."""
    d = find_final_decision(pred)
    evidence = first_present(d, ["E", "evidence", "evidence_trace", "trace"])
    if isinstance(evidence, dict):
        reasoning = first_present(evidence, ["reasoning_path", "path", "steps"])
        explanation = first_present(evidence, ["explanation", "summary", "rationale"])
    else:
        reasoning = evidence
        explanation = first_present(d, ["explanation", "rationale", "reason"])
    sources = first_present(d, ["S", "sources", "citations", "evidence_sources"])

    return [
        ("Recommend", as_text(first_present(d, ["M_rec", "recommend", "recommended", "recommendations"]))),
        ("Avoid", as_text(first_present(d, ["M_avoid", "avoid", "avoidance"]))),
        ("Caution and monitoring", as_text(first_present(d, ["M_caution", "caution", "monitoring"]))),
        ("Safer alternatives", as_text(first_present(d, ["M_alt", "alternative", "alternatives"]))),
        ("Risk level", as_text(first_present(d, ["M_level", "risk_level", "severity"]))),
        ("Evidence path", as_text(reasoning)),
        ("Explanation", as_text(explanation)),
        ("Sources", as_text(sources)),
    ]


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, size: float = 8.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_doc_defaults(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(0.6)
    sec.bottom_margin = Inches(0.6)
    sec.left_margin = Inches(0.65)
    sec.right_margin = Inches(0.65)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(9)
    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        styles[style_name].font.name = "Times New Roman"


def add_title(doc: Document, reviewer_id: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ATLAS Blinded Clinician Review")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(15)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"Reviewer {reviewer_id}")
    r2.bold = True
    r2.font.size = Pt(11)

    doc.add_paragraph(
        "Review each case using the full patient information. Compare Output A and Output B without inferring the model identity. "
        "Score both outputs before selecting a preference. Do not use external automated scores or reference labels during this pass."
    )


def add_rubric(doc: Document) -> None:
    doc.add_heading("Rating rubric", level=1)
    t = doc.add_table(rows=1, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    headers = ["Criterion", "Definition", "Scale anchors"]
    for i, h in enumerate(headers):
        set_cell_text(t.rows[0].cells[i], h, bold=True)
        shade_cell(t.rows[0].cells[i], "D9EAD3")
    for name, definition in CRITERIA:
        row = t.add_row().cells
        set_cell_text(row[0], name, bold=True)
        set_cell_text(row[1], definition)
        set_cell_text(row[2], "1 = major errors\n2 = substantial limitations\n3 = acceptable with gaps\n4 = strong\n5 = complete and well supported")
    doc.add_paragraph(
        "Potentially unsafe recommendation: mark Yes when the output recommends, permits, or fails to exclude an option that could cause clinically meaningful harm under the full patient state."
    )
    doc.add_paragraph(
        "Reference audit: complete the separate second pass only after you submit the primary ratings. The second pass may show the guideline-derived reference and source excerpt."
    )


def add_case_section(
    doc: Document,
    packet_no: int,
    dataset: str,
    case_obj: dict[str, Any],
    output_a: dict[str, Any],
    output_b: dict[str, Any],
) -> None:
    if packet_no > 1:
        doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading(f"Case {packet_no:02d}", level=1)
    p = doc.add_paragraph()
    p.add_run("Benchmark: ").bold = True
    p.add_run(dataset)

    doc.add_heading("Full patient information", level=2)
    clean_case = sanitize_case(deepcopy(case_obj))
    patient_text = json.dumps(clean_case, ensure_ascii=False, indent=2)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(patient_text)
    r.font.name = "Consolas"
    r.font.size = Pt(7.5)

    for label, pred in (("Output A", output_a), ("Output B", output_b)):
        doc.add_heading(label, level=2)
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for field, value in normalized_output(pred):
            cells = table.add_row().cells
            set_cell_text(cells[0], field, bold=True)
            shade_cell(cells[0], "EDEDED")
            set_cell_text(cells[1], value, size=8)

    doc.add_heading("Primary ratings", level=2)
    rt = doc.add_table(rows=1, cols=4)
    rt.style = "Table Grid"
    rt.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Criterion", "Output A (1–5)", "Output B (1–5)", "Comments"]):
        set_cell_text(rt.rows[0].cells[i], h, bold=True)
        shade_cell(rt.rows[0].cells[i], "D9EAD3")
    for name, _ in CRITERIA:
        cells = rt.add_row().cells
        set_cell_text(cells[0], name)
        set_cell_text(cells[1], "")
        set_cell_text(cells[2], "")
        set_cell_text(cells[3], "")

    q = doc.add_table(rows=3, cols=2)
    q.style = "Table Grid"
    prompts = [
        "Potentially unsafe recommendation",
        "Preferred output",
        "Optional clinical comment",
    ]
    values = [
        "Output A: Yes / No     Output B: Yes / No",
        "A / B / Tie",
        "",
    ]
    for i in range(3):
        set_cell_text(q.rows[i].cells[0], prompts[i], bold=True)
        set_cell_text(q.rows[i].cells[1], values[i])


def write_rating_template(path: Path, reviewer_id: str, packet_rows: list[dict[str, Any]]) -> None:
    fields = [
        "reviewer_id",
        "packet_case",
        "dataset",
        "case_code",
        "A_correctness",
        "A_safety",
        "A_completeness",
        "A_actionability",
        "A_evidence_consistency",
        "A_unsafe",
        "B_correctness",
        "B_safety",
        "B_completeness",
        "B_actionability",
        "B_evidence_consistency",
        "B_unsafe",
        "preference",
        "reference_issue_pass2",
        "comments",
    ]
    with path.open("w", newline="", encoding="utf-8-sig") as f:
        w = csv.DictWriter(f, fieldnames=fields)
        w.writeheader()
        for row in packet_rows:
            w.writerow(
                {
                    "reviewer_id": reviewer_id,
                    "packet_case": row["packet_case"],
                    "dataset": row["dataset"],
                    "case_code": f"{reviewer_id}-{int(row['packet_case']):02d}",
                }
            )


def build_inputs(args: argparse.Namespace) -> dict[str, dict[str, dict[str, Any]]]:
    data: dict[str, dict[str, dict[str, Any]]] = {}
    specs = {
        "Western": (args.western_cases, args.western_atlas, args.western_gemini),
        "GeriMedBench": (args.geri_cases, args.geri_atlas, args.geri_gemini),
    }
    for dataset, (cases_p, atlas_p, gemini_p) in specs.items():
        cases = index_records(load_records(cases_p), args.case_id_key)
        atlas = index_records(load_records(atlas_p), args.case_id_key)
        gemini = index_records(load_records(gemini_p), args.case_id_key)
        common = sorted(set(cases) & set(atlas) & set(gemini))
        if len(common) < args.n_per_benchmark:
            raise ValueError(
                f"{dataset} has only {len(common)} matched cases. Need {args.n_per_benchmark}."
            )
        data[dataset] = {
            cid: {"case": cases[cid], MODEL_ATLAS: atlas[cid], MODEL_GEMINI: gemini[cid]}
            for cid in common
        }
    return data


def main() -> int:
    args = parse_args()
    out_dir = Path(args.out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    packets_dir = out_dir / "review_packets"
    ratings_dir = out_dir / "rating_templates"
    packets_dir.mkdir(exist_ok=True)
    ratings_dir.mkdir(exist_ok=True)

    data = build_inputs(args)
    sample_rng = random.Random(args.seed)
    selected: list[tuple[str, str]] = []
    for dataset in ("Western", "GeriMedBench"):
        ids = sorted(data[dataset])
        chosen = sample_rng.sample(ids, args.n_per_benchmark)
        selected.extend((dataset, cid) for cid in chosen)

    manifest_path = out_dir / "sample_manifest.csv"
    with manifest_path.open("w", newline="", encoding="utf-8-sig") as f:
        w = csv.DictWriter(f, fieldnames=["dataset", "case_id", "selection_seed"])
        w.writeheader()
        for dataset, cid in selected:
            w.writerow({"dataset": dataset, "case_id": cid, "selection_seed": args.seed})

    key_rows: list[dict[str, Any]] = []
    for reviewer_idx in range(1, args.reviewers + 1):
        reviewer_id = f"R{reviewer_idx}"
        rng = random.Random(args.seed + reviewer_idx * 1009)
        order = selected[:]
        rng.shuffle(order)
        doc = Document()
        set_doc_defaults(doc)
        add_title(doc, reviewer_id)
        add_rubric(doc)
        packet_rows: list[dict[str, Any]] = []
        for packet_no, (dataset, cid) in enumerate(order, start=1):
            if rng.random() < 0.5:
                a_method, b_method = MODEL_ATLAS, MODEL_GEMINI
            else:
                a_method, b_method = MODEL_GEMINI, MODEL_ATLAS
            item = data[dataset][cid]
            add_case_section(
                doc,
                packet_no,
                dataset,
                item["case"],
                item[a_method],
                item[b_method],
            )
            row = {
                "reviewer_id": reviewer_id,
                "packet_case": packet_no,
                "dataset": dataset,
                "case_id": cid,
                "A_method": a_method,
                "B_method": b_method,
            }
            packet_rows.append(row)
            key_rows.append(row)
        doc.save(packets_dir / f"reviewer_{reviewer_id}_packet.docx")
        write_rating_template(ratings_dir / f"reviewer_{reviewer_id}_ratings.csv", reviewer_id, packet_rows)

    key_path = out_dir / "PRIVATE_randomization_key.csv"
    with key_path.open("w", newline="", encoding="utf-8-sig") as f:
        fields = ["reviewer_id", "packet_case", "dataset", "case_id", "A_method", "B_method"]
        w = csv.DictWriter(f, fieldnames=fields)
        w.writeheader()
        w.writerows(key_rows)

    readme = out_dir / "PREPARATION_COMPLETE.txt"
    readme.write_text(
        "Preparation completed. Keep PRIVATE_randomization_key.csv away from reviewers. "
        "Send each reviewer only the matching DOCX packet and CSV rating template.\n",
        encoding="utf-8",
    )
    print(f"[OK] Prepared {len(selected)} cases for {args.reviewers} reviewers in {out_dir}")
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:
        print(f"[ERROR] {exc}", file=sys.stderr)
        raise
