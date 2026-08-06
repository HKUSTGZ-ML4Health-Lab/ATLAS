#!/usr/bin/env python3
# AAAI-27 paper reference: Blinded Clinician Evaluation and Table 2.
"""Aggregate blinded clinician-review ratings and create paper-ready outputs.

The script maps A/B labels back to ATLAS and Gemini through the private key,
checks all ratings, computes five mean clinician scores, case-level majority
unsafe and preference outcomes, ordinal Krippendorff alpha, bootstrap
confidence intervals, and paper-ready DOCX, CSV, JSON, and LaTeX files.
"""

from __future__ import annotations

import argparse
import csv
import json
import math
import random
import re
import statistics
import sys
from collections import Counter, defaultdict
from pathlib import Path
from typing import Any, Iterable

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from scipy.stats import wilcoxon

MODELS = ["Gemini 3.1 Pro Preview", "ATLAS"]
CRITERIA = [
    ("correctness", "Correct."),
    ("safety", "Safety"),
    ("completeness", "Complete."),
    ("actionability", "Action."),
    ("evidence_consistency", "Evidence"),
]


def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(description=__doc__)
    p.add_argument("--ratings-dir", required=True)
    p.add_argument("--key", required=True)
    p.add_argument("--out-dir", required=True)
    p.add_argument("--bootstrap", type=int, default=10000)
    p.add_argument("--seed", type=int, default=20260729)
    return p.parse_args()


def read_csvs(ratings_dir: Path) -> pd.DataFrame:
    files = sorted(ratings_dir.glob("*.csv"))
    if not files:
        raise FileNotFoundError(f"No CSV files found in {ratings_dir}")
    frames = [pd.read_csv(p, dtype=str, keep_default_na=False) for p in files]
    df = pd.concat(frames, ignore_index=True)
    required = {
        "reviewer_id",
        "packet_case",
        "A_correctness",
        "A_safety",
        "A_completeness",
        "A_actionability",
        "A_evidence_consistency",
        "A_unsafe",
        "B_correctness",
        "B_safety",
        "B_completeness",
        "B_actionability",
        "B_evidence_consistency",
        "B_unsafe",
        "preference",
        "reference_issue_pass2",
    }
    missing = required - set(df.columns)
    if missing:
        raise ValueError(f"Ratings files miss columns: {sorted(missing)}")
    df["packet_case"] = df["packet_case"].astype(str).str.strip()
    return df


def normalize_yes_no(value: str, field: str) -> int:
    s = str(value).strip().lower()
    if s in {"yes", "y", "1", "true"}:
        return 1
    if s in {"no", "n", "0", "false"}:
        return 0
    raise ValueError(f"{field} must be Yes or No, got {value!r}")


def normalize_preference(value: str) -> str:
    s = str(value).strip().lower()
    if s in {"a", "output a"}:
        return "A"
    if s in {"b", "output b"}:
        return "B"
    if s in {"tie", "equal", "same"}:
        return "Tie"
    raise ValueError(f"preference must be A, B, or Tie, got {value!r}")


def parse_score(value: str, field: str) -> float:
    try:
        x = float(str(value).strip())
    except ValueError as exc:
        raise ValueError(f"{field} must be a score from 1 to 5, got {value!r}") from exc
    if x < 1 or x > 5 or not x.is_integer():
        raise ValueError(f"{field} must be an integer from 1 to 5, got {value!r}")
    return x


def load_key(path: Path) -> pd.DataFrame:
    key = pd.read_csv(path, dtype=str, keep_default_na=False)
    required = {"reviewer_id", "packet_case", "dataset", "case_id", "A_method", "B_method"}
    missing = required - set(key.columns)
    if missing:
        raise ValueError(f"Key file misses columns: {sorted(missing)}")
    key["packet_case"] = key["packet_case"].astype(str).str.strip()
    return key


def validate_and_longify(ratings: pd.DataFrame, key: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame]:
    merged = ratings.merge(
        key,
        on=["reviewer_id", "packet_case"],
        how="inner",
        validate="one_to_one",
        suffixes=("", "_key"),
    )
    if len(merged) != len(ratings) or len(merged) != len(key):
        raise ValueError(
            f"Ratings/key mismatch. ratings={len(ratings)}, key={len(key)}, merged={len(merged)}"
        )

    long_rows: list[dict[str, Any]] = []
    pref_rows: list[dict[str, Any]] = []
    for _, row in merged.iterrows():
        for side in ("A", "B"):
            model = row[f"{side}_method"]
            if model not in MODELS:
                raise ValueError(f"Unexpected model name in key: {model}")
            out = {
                "reviewer_id": row["reviewer_id"],
                "packet_case": row["packet_case"],
                "dataset": row["dataset_key"] if "dataset_key" in row and row["dataset_key"] else row.get("dataset", ""),
                "case_id": row["case_id"],
                "unit_id": f"{row['case_id']}::{model}",
                "model": model,
                "unsafe": normalize_yes_no(row[f"{side}_unsafe"], f"{side}_unsafe"),
            }
            for criterion, _ in CRITERIA:
                out[criterion] = parse_score(row[f"{side}_{criterion}"], f"{side}_{criterion}")
            long_rows.append(out)

        preference = normalize_preference(row["preference"])
        preferred_model = "Tie" if preference == "Tie" else row[f"{preference}_method"]
        pref_rows.append(
            {
                "reviewer_id": row["reviewer_id"],
                "dataset": row["dataset_key"] if "dataset_key" in row and row["dataset_key"] else row.get("dataset", ""),
                "case_id": row["case_id"],
                "preferred_model": preferred_model,
                "reference_issue": normalize_yes_no(row["reference_issue_pass2"], "reference_issue_pass2"),
                "comments": row.get("comments", ""),
            }
        )
    return pd.DataFrame(long_rows), pd.DataFrame(pref_rows)


def coincidence_alpha(data: pd.DataFrame, value_col: str, ordinal: bool = True) -> float:
    """Compute Krippendorff alpha from reviewer x unit ratings.

    For ordinal data, this uses Krippendorff's marginal-frequency distance.
    For binary or nominal data, set ordinal=False.
    """
    units: list[list[Any]] = []
    for _, group in data.groupby("unit_id", sort=False):
        vals = [v for v in group[value_col].tolist() if pd.notna(v)]
        if len(vals) >= 2:
            units.append(vals)
    if not units:
        return float("nan")

    categories = sorted({v for vals in units for v in vals})
    cat_index = {c: i for i, c in enumerate(categories)}
    k = len(categories)
    coincidence = np.zeros((k, k), dtype=float)

    for vals in units:
        m = len(vals)
        for i, vi in enumerate(vals):
            for j, vj in enumerate(vals):
                if i == j:
                    continue
                coincidence[cat_index[vi], cat_index[vj]] += 1.0 / (m - 1)

    marginals = coincidence.sum(axis=1)
    n = marginals.sum()
    if n <= 1:
        return float("nan")

    expected = np.zeros_like(coincidence)
    for i in range(k):
        for j in range(k):
            if i == j:
                expected[i, j] = marginals[i] * (marginals[i] - 1) / (n - 1)
            else:
                expected[i, j] = marginals[i] * marginals[j] / (n - 1)

    def distance(i: int, j: int) -> float:
        if i == j:
            return 0.0
        if not ordinal:
            return 1.0
        lo, hi = sorted((i, j))
        span = marginals[lo : hi + 1].sum() - (marginals[lo] + marginals[hi]) / 2.0
        return float(span * span)

    do = 0.0
    de = 0.0
    for i in range(k):
        for j in range(k):
            d = distance(i, j)
            do += coincidence[i, j] * d
            de += expected[i, j] * d
    do /= n
    de /= n
    if de == 0:
        return 1.0 if do == 0 else float("nan")
    return 1.0 - do / de


def case_level_scores(long_df: pd.DataFrame) -> pd.DataFrame:
    return (
        long_df.groupby(["case_id", "dataset", "model"], as_index=False)[[c for c, _ in CRITERIA]]
        .mean()
    )


def bootstrap_ci(case_df: pd.DataFrame, model: str, criterion: str, n_boot: int, seed: int) -> tuple[float, float]:
    vals = case_df.loc[case_df["model"] == model, criterion].to_numpy(dtype=float)
    if len(vals) == 0:
        return float("nan"), float("nan")
    rng = np.random.default_rng(seed)
    idx = rng.integers(0, len(vals), size=(n_boot, len(vals)))
    means = vals[idx].mean(axis=1)
    return float(np.quantile(means, 0.025)), float(np.quantile(means, 0.975))


def majority_label(values: Iterable[str], valid_labels: set[str]) -> str:
    counts = Counter(v for v in values if v in valid_labels)
    if not counts:
        return "Tie"
    top = counts.most_common()
    if len(top) == 1:
        return top[0][0]
    if top[0][1] == top[1][1]:
        return "Tie"
    return top[0][0]


def majority_binary(values: Iterable[int]) -> int | None:
    vals = [int(v) for v in values]
    if not vals:
        return None
    yes = sum(vals)
    no = len(vals) - yes
    if yes == no:
        return None
    return 1 if yes > no else 0


def holm_adjust(p_values: dict[str, float]) -> dict[str, float]:
    items = sorted(p_values.items(), key=lambda kv: kv[1])
    m = len(items)
    adjusted: dict[str, float] = {}
    prev = 0.0
    for rank, (name, p) in enumerate(items, start=1):
        val = min(1.0, (m - rank + 1) * p)
        val = max(prev, val)
        adjusted[name] = val
        prev = val
    return adjusted


def compute_summary(long_df: pd.DataFrame, pref_df: pd.DataFrame, args: argparse.Namespace) -> dict[str, Any]:
    case_df = case_level_scores(long_df)
    summary: dict[str, Any] = {"methods": {}, "agreement": {}, "tests": {}, "preference": {}, "reference_audit": {}}

    for model in MODELS:
        model_data = long_df[long_df["model"] == model]
        method_summary: dict[str, Any] = {}
        for i, (criterion, _) in enumerate(CRITERIA):
            mean = float(model_data[criterion].astype(float).mean())
            sd = float(model_data[criterion].astype(float).std(ddof=1))
            lo, hi = bootstrap_ci(case_df, model, criterion, args.bootstrap, args.seed + i * 31 + MODELS.index(model) * 1000)
            method_summary[criterion] = {"mean": mean, "sd": sd, "ci95": [lo, hi]}

        unsafe_cases = 0
        total_cases = 0
        for _, group in model_data.groupby("case_id"):
            maj = majority_binary(group["unsafe"].tolist())
            if maj is not None:
                total_cases += 1
                unsafe_cases += maj
        method_summary["unsafe_cases"] = {
            "n": int(unsafe_cases),
            "denominator": int(total_cases),
            "percent": 100.0 * unsafe_cases / total_cases if total_cases else float("nan"),
        }
        summary["methods"][model] = method_summary

    for criterion, _ in CRITERIA:
        summary["agreement"][criterion] = coincidence_alpha(long_df, criterion, ordinal=True)

    # Paired case-level Wilcoxon tests with Holm correction.
    raw_p: dict[str, float] = {}
    for criterion, _ in CRITERIA:
        pivot = case_df.pivot(index="case_id", columns="model", values=criterion).dropna()
        if set(MODELS).issubset(pivot.columns) and len(pivot) > 0:
            diff = pivot["ATLAS"] - pivot["Gemini 3.1 Pro Preview"]
            if np.allclose(diff.to_numpy(), 0):
                p = 1.0
            else:
                p = float(wilcoxon(pivot["ATLAS"], pivot["Gemini 3.1 Pro Preview"], zero_method="wilcox").pvalue)
            raw_p[criterion] = p
            summary["tests"][criterion] = {
                "mean_difference": float(diff.mean()),
                "wilcoxon_p": p,
            }
    adjusted = holm_adjust(raw_p)
    for criterion, p in adjusted.items():
        summary["tests"][criterion]["holm_p"] = p

    # Case-level majority preference.
    case_pref: dict[str, str] = {}
    for case_id, group in pref_df.groupby("case_id"):
        case_pref[str(case_id)] = majority_label(group["preferred_model"], set(MODELS) | {"Tie"})
    pref_counts = Counter(case_pref.values())
    total_pref = len(case_pref)
    summary["preference"] = {
        model: {
            "n": int(pref_counts.get(model, 0)),
            "denominator": int(total_pref),
            "percent": 100.0 * pref_counts.get(model, 0) / total_pref if total_pref else float("nan"),
        }
        for model in MODELS
    }
    summary["preference"]["Tie"] = {
        "n": int(pref_counts.get("Tie", 0)),
        "denominator": int(total_pref),
        "percent": 100.0 * pref_counts.get("Tie", 0) / total_pref if total_pref else float("nan"),
    }

    # Reference audit majority by case.
    flagged = 0
    ref_total = 0
    for _, group in pref_df.groupby("case_id"):
        maj = majority_binary(group["reference_issue"].tolist())
        if maj is not None:
            ref_total += 1
            flagged += maj
    summary["reference_audit"] = {
        "flagged_cases": int(flagged),
        "denominator": int(ref_total),
        "percent": 100.0 * flagged / ref_total if ref_total else float("nan"),
    }
    return summary


def fmt_score(value: float) -> str:
    return f"{value:.2f}"


def fmt_n_pct(obj: dict[str, Any]) -> str:
    return f"{obj['n']}/{obj['denominator']} ({obj['percent']:.1f}\\%)"


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell(cell, text: str, bold: bool = False, size: float = 8.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def create_docx(summary: dict[str, Any], out_path: Path) -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.6)
    sec.right_margin = Inches(0.6)
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Blinded Clinician Evaluation Results")
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = "Times New Roman"

    caption = (
        "Table. Blinded clinician evaluation on a benchmark-stratified sample of multimorbidity cases. "
        "Scores use a five-point scale. Unsafe and preferred values report case-level majority judgments."
    )
    cp = doc.add_paragraph(caption)
    cp.alignment = WD_ALIGN_PARAGRAPH.LEFT

    table = doc.add_table(rows=1, cols=8)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Method", "Correct.", "Safety", "Complete.", "Action.", "Evidence", "Unsafe cases", "Preferred cases"]
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, bold=True)
        shade_cell(table.rows[0].cells[i], "D9EAD3")

    for model in MODELS:
        cells = table.add_row().cells
        set_cell(cells[0], model, bold=(model == "ATLAS"))
        for j, (criterion, _) in enumerate(CRITERIA, start=1):
            set_cell(cells[j], fmt_score(summary["methods"][model][criterion]["mean"]), bold=(model == "ATLAS"))
        set_cell(cells[6], fmt_n_pct(summary["methods"][model]["unsafe_cases"]), bold=(model == "ATLAS"))
        set_cell(cells[7], fmt_n_pct(summary["preference"][model]), bold=(model == "ATLAS"))

    note = doc.add_paragraph()
    note.add_run("Note. ").bold = True
    alphas = [summary["agreement"][criterion] for criterion, _ in CRITERIA]
    note.add_run(
        f"Ordinal Krippendorff's alpha ranges from {min(alphas):.2f} to {max(alphas):.2f} across the five rating criteria. "
        f"Reviewers report ties in {summary['preference']['Tie']['n']} cases. "
        f"The separate reference audit flags {summary['reference_audit']['flagged_cases']} cases by majority judgment."
    )
    doc.save(out_path)


def create_latex(summary: dict[str, Any], out_path: Path) -> None:
    lines = [
        r"\begin{table*}[t]",
        r"\centering",
        r"\caption{Blinded clinician evaluation on a benchmark-stratified sample of multimorbidity cases. Scores use a five-point scale. Unsafe and preferred values report case-level majority judgments.}",
        r"\label{tab:expert_review}",
        r"\small",
        r"\begin{tabular}{lccccccc}",
        r"\toprule",
        r"Method & Correct. & Safety & Complete. & Action. & Evidence & Unsafe cases & Preferred cases \\",
        r"\midrule",
    ]
    for model in MODELS:
        vals = [fmt_score(summary["methods"][model][criterion]["mean"]) for criterion, _ in CRITERIA]
        unsafe = fmt_n_pct(summary["methods"][model]["unsafe_cases"])
        preferred = fmt_n_pct(summary["preference"][model])
        name = r"\textbf{ATLAS}" if model == "ATLAS" else model
        lines.append(f"{name} & " + " & ".join(vals + [unsafe, preferred]) + r" \\")
    lines += [
        r"\bottomrule",
        r"\end{tabular}",
    ]
    alphas = [summary["agreement"][criterion] for criterion, _ in CRITERIA]
    lines.append(
        r"\vspace{1mm}\parbox{0.98\textwidth}{\footnotesize "
        + f"Ordinal Krippendorff's $\\alpha$ ranges from {min(alphas):.2f} to {max(alphas):.2f} across the five criteria. "
        + f"Reviewers report ties in {summary['preference']['Tie']['n']} cases. "
        + f"The separate reference audit flags {summary['reference_audit']['flagged_cases']} cases by majority judgment."
        + r"}"
    )
    lines.append(r"\end{table*}")
    out_path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def supported_dimensions(summary: dict[str, Any]) -> list[str]:
    supported: list[str] = []
    for criterion, label in CRITERIA:
        atlas = summary["methods"]["ATLAS"][criterion]["mean"]
        gemini = summary["methods"]["Gemini 3.1 Pro Preview"][criterion]["mean"]
        if atlas > gemini:
            supported.append(label.rstrip("."))
    return supported


def create_result_text(summary: dict[str, Any], out_path: Path) -> None:
    dims = supported_dimensions(summary)
    if len(dims) == 0:
        dim_text = "no rating dimension"
    elif len(dims) == 1:
        dim_text = dims[0]
    else:
        dim_text = ", ".join(dims[:-1]) + ", and " + dims[-1]
    atlas_pref = summary["preference"]["ATLAS"]["n"]
    gem_pref = summary["preference"]["Gemini 3.1 Pro Preview"]["n"]
    ties = summary["preference"]["Tie"]["n"]
    atlas_unsafe = summary["methods"]["ATLAS"]["unsafe_cases"]["n"]
    gem_unsafe = summary["methods"]["Gemini 3.1 Pro Preview"]["unsafe_cases"]["n"]
    alphas = [summary["agreement"][criterion] for criterion, _ in CRITERIA]
    ref = summary["reference_audit"]["flagged_cases"]

    text = rf"""\subsection{{Blinded Expert Evaluation}}

Table~\ref{{tab:expert_review}} reports the blinded clinician evaluation.
Clinicians assign ATLAS higher mean scores for {dim_text}.
Case-level majority judgments prefer ATLAS in {atlas_pref} cases, prefer
Gemini 3.1 Pro Preview in {gem_pref} cases, and report ties in {ties} cases.
Reviewers flag potentially unsafe recommendations in {atlas_unsafe} ATLAS
cases and {gem_unsafe} Gemini cases. Ordinal Krippendorff's $\alpha$ ranges
from {min(alphas):.2f} to {max(alphas):.2f} across the five criteria. The
separate reference audit identifies {ref} cases with majority disagreement.
These findings support the clinician-rated clinical quality of the outputs and
do not establish real-world clinical effectiveness.
"""
    out_path.write_text(text, encoding="utf-8")


def main() -> int:
    args = parse_args()
    ratings_dir = Path(args.ratings_dir)
    key_path = Path(args.key)
    out_dir = Path(args.out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    ratings = read_csvs(ratings_dir)
    key = load_key(key_path)
    long_df, pref_df = validate_and_longify(ratings, key)
    summary = compute_summary(long_df, pref_df, args)

    long_df.to_csv(out_dir / "expert_review_long.csv", index=False, encoding="utf-8-sig")
    pref_df.to_csv(out_dir / "expert_review_preferences.csv", index=False, encoding="utf-8-sig")
    case_level_scores(long_df).to_csv(out_dir / "expert_review_case_means.csv", index=False, encoding="utf-8-sig")
    (out_dir / "expert_review_results.json").write_text(
        json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    rows = []
    for model in MODELS:
        row: dict[str, Any] = {"Method": model}
        for criterion, label in CRITERIA:
            row[label] = summary["methods"][model][criterion]["mean"]
        row["Unsafe cases"] = fmt_n_pct(summary["methods"][model]["unsafe_cases"])
        row["Preferred cases"] = fmt_n_pct(summary["preference"][model])
        rows.append(row)
    pd.DataFrame(rows).to_csv(out_dir / "expert_review_summary_table.csv", index=False, encoding="utf-8-sig")

    create_docx(summary, out_dir / "expert_review_results_table.docx")
    create_latex(summary, out_dir / "expert_review_results_table.tex")
    create_result_text(summary, out_dir / "expert_review_results_paragraph.tex")
    print(f"[OK] Clinician-review outputs written to {out_dir}")
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:
        print(f"[ERROR] {exc}", file=sys.stderr)
        raise
